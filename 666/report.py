from docx import Document
from openpyxl import Workbook

def save_doc(filename, data):
    doc = Document()
    doc.add_heading("Отчёт по энергии", 0)

    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.save(filename)

def save_xls(filename, data):
    wb = Workbook()
    ws = wb.active

    row = 1
    for key, value in data.items():
        ws.cell(row=row, column=1, value=key)
        ws.cell(row=row, column=2, value=str(value))
        row += 1

    wb.save(filename)
