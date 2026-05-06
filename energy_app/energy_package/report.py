from docx import Document
from openpyxl import Workbook

def save_doc(filename, data):
    doc = Document()
    doc.add_heading("Отчёт по энергии", 0)

    for k, v in data.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.save(filename)

def save_xls(filename, data):
    wb = Workbook()
    ws = wb.active

    row = 1
    for k, v in data.items():
        ws.cell(row=row, column=1, value=k)
        ws.cell(row=row, column=2, value=str(v))
        row += 1

    wb.save(filename)
