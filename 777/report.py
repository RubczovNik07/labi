from docx import Document
from openpyxl import Workbook


class Report:

    def __init__(self, data):
        self.data = data

    def save_docx(self, filename):
        doc = Document()
        doc.add_heading("Отчёт по энергии", 0)

        for k, v in self.data.items():
            doc.add_paragraph(f"{k}: {v}")

        doc.save(filename)

    def save_xlsx(self, filename):
        wb = Workbook()
        ws = wb.active

        row = 1
        for k, v in self.data.items():
            ws.cell(row=row, column=1, value=k)
            ws.cell(row=row, column=2, value=str(v))
            row += 1

        wb.save(filename)

    # dunder
    def len(self):
        return len(self.data)

    def str(self):
        return "Report"
